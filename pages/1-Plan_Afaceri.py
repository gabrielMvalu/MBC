# pages/Plan_faceri.py
import streamlit as st
import pandas as pd
from docx import Document
from pages.constatator import extrage_informatii_firma, extrage_asociati_admini, extrage_situatie_angajati, extrage_coduri_caen
from pages.datesolicitate import extrage_date_solicitate
from pages.bilantsianaliza import extrage_date_bilant, extrage_date_contpp, extrage_indicatori_financiari
from pages.ssiutilaje import extrage_pozitii, coreleaza_date

st.set_page_config(layout="wide")

st.header(':blue[Completare Document cu Placeholder-uri]', divider='rainbow')

st.write("Selectează codul CAEN pentru activitatea ta:")
caen_options = {
    "CAEN 4312": "Lucrări de pregătire a terenului",
    "CAEN 4211": "Lucrări de construcții a drumurilor și autostrăzilor",
    "CAEN 4399": "Alte lucrări speciale de construcții n.c.a.",
    "CAEN 3832": "Recuperarea materialelor reciclabile sortate"
}

for caen_code, description in caen_options.items():
    if st.checkbox(f"{caen_code} - {description}"):
        st.session_state.codCAEN = caen_code.split()[1]


col1, col2 = st.columns(2)
with col1:
    uploaded_template = st.file_uploader("Încărcați documentul", type=["docx"], key="template")
    st.info('Adaugati macheta Planului de Afaceri sau a Cererii de Finantare', icon="⬆️")

with col2:
    uploaded_document = st.file_uploader("Încărcați documentul", type=["docx"], key="document")
    st.info('Adaugati Raport interogare', icon="⬆️")

st.divider()

col3, col4 = st.columns(2)
with col3:
    st.success('Adaugati fisierul excel Date Solicitate', icon="⬇️")
    uploaded_file1 = st.file_uploader("Încărcați documentul", type=["xlsx"], key="excelSolicitate")
    
with col4:
    st.success('Adaugati fisierul excel Anexa 3 Macheta financiara', icon="⬇️")
    uploaded_file2 = st.file_uploader("Încărcați documentul", type=["xlsx"], key="excelBCAP")
    

        
    

if uploaded_template is not None and uploaded_document is not None and uploaded_file1 is not None and uploaded_file2 is not None:
    template_doc = Document(uploaded_template)
    st.toast('Incepem procesarea Planului de afaceri', icon='⭐')     
    constatator_doc = Document(uploaded_document)
    date_solicitate_doc = pd.read_excel(uploaded_file1, sheet_name='date solicitate')
    df = pd.read_excel(uploaded_file2, sheet_name='1-Bilant')
    df1 = pd.read_excel(uploaded_file2, sheet_name='2-ContPP')
    df2 = pd.read_excel(uploaded_file2, sheet_name='1D-Analiza_fin_indicatori')    
    
    df_financiar = pd.read_excel(uploaded_file2, sheet_name='P. FINANCIAR')
    date_financiare = extrage_pozitii(df_financiar)
    if 'codCAEN' in st.session_state and date_financiare:
        rezultate_corelate, rezultate_corelate1, rezultate_corelate2 = coreleaza_date(date_financiare)
        rezultate_text = '\n'.join([rezultat for _, _, rezultat in rezultate_corelate])
        cheltuieli_text = '\n'.join([rezultat for _, _, rezultat in rezultate_corelate1])
        cantitati_corelate = [pd.to_numeric(item[1], errors='coerce') for item in rezultate_corelate]
        cantitati_corelate = [0 if pd.isna(x) else x for x in cantitati_corelate]
        numar_total_utilaje = sum(cantitati_corelate)
        rezultate_corelate, rezultate_corelate1, rezultate_corelate2 = coreleaza_date(date_financiare)
        rezultate2_text = '\n'.join([f"{descriere}" for nume, _, descriere in rezultate_corelate2])


    
    informatii_firma = extrage_informatii_firma(constatator_doc)
    asociati_info, administratori_info = extrage_asociati_admini(constatator_doc)
    situatie_angajati = extrage_situatie_angajati(constatator_doc)
    full_text_constatator = "\n".join([p.text for p in constatator_doc.paragraphs])
    coduri_caen = extrage_coduri_caen(full_text_constatator)
    date_solicitate = extrage_date_solicitate(date_solicitate_doc)
    capital_propriu = extrage_date_bilant(df)
    cifra_venit_rezultat = extrage_date_contpp(df1)
    rata_rent_grad = extrage_indicatori_financiari(df2)

    def curata_duplicate_coduri_caen(coduri_caen):
        coduri_unice = {}
        for cod, descriere in coduri_caen:
            coduri_unice[cod] = descriere
        return list(coduri_unice.items())

    coduri_caen_curatate = curata_duplicate_coduri_caen(coduri_caen)
    
    adrese_secundare_text = '\n'.join(informatii_firma.get('Adresa sediul secundar', [])) if informatii_firma.get('Adresa sediul secundar', []) else "N/A"
    asociati_text = '\n'.join(asociati_info) if asociati_info else "N/A"
    administratori_text = administratori_info if administratori_info else "N/A"
    coduri_caen_text = '\n'.join([f"{cod} - {descriere}" for cod, descriere in coduri_caen_curatate]) if coduri_caen_curatate else "N/A"    

    placeholders = {
        "#SRL": str(informatii_firma.get('Denumirea firmei', 'N/A')),
        "#CUI": str(informatii_firma.get('Codul unic de înregistrare (CUI)', 'N/A')),
        "#Nr_inmatriculare": str(informatii_firma.get('Numărul de ordine în Registrul Comerțului', 'N/A')),
        "#data_infiintare": str(informatii_firma.get('Data înființării', 'N/A')),
        "#Adresa_sediu": str(informatii_firma.get('Adresa sediului social', 'N/A')),
        "#Adresa_pct_lucru": str(adrese_secundare_text),
        "#Asociati": str(asociati_text),
        "#Administrator": str(administratori_text),
        "#activitatePrincipala": str(informatii_firma.get('Activitate principală', 'N/A')),
        "#CAENautorizate": str(coduri_caen_text),
        "#categ_intreprindere": str(date_solicitate.get('Categorie întreprindere', 'N/A')),
        "#Firme_legate": str(date_solicitate.get('Firme legate', 'N/A')),
        "#Tip_investitie": str(date_solicitate.get('Tipul investiției', 'N/A')),
        "#activitate": str(date_solicitate.get('Activitate', 'N/A')),
        "#CAEN": str(date_solicitate.get('Cod CAEN', 'N/A')),
        "#nr_locuri_munca_noi": str(date_solicitate.get('Număr locuri de muncă noi', 'N/A')),
        "#Judet": str(date_solicitate.get('Județ', 'N/A')),

        "#Utilaj": str(rezultate_text),
        "#cheltuieli_proiect_din_buget_excel": str(cheltuieli_text),
        "#DescriereUtilaje" : str(rezultate2_text),
        "#nr_utilaje": str(numar_total_utilaje),
        
        "#utilaj_dizabilitati": str(date_solicitate.get('Utilaj pentru persoane cu dizabilități', 'N/A')),
        "#utilaj_cu_tocator": str(date_solicitate.get('Utilaj cu tocător', 'N/A')),
        "#adresa_loc_implementare": str(date_solicitate.get('Adresa locației de implementare', 'N/A')),
        "#nrClasareNotificare": str(date_solicitate.get('Număr clasare notificare', 'N/A')),
        "#clientiActuali": str(date_solicitate.get('Clienți actuali', 'N/A')),
        "#furnizori": str(date_solicitate.get('Furnizori', 'N/A')),
        "#tip_activitate": str(date_solicitate.get('Tip activitate', 'N/A')),
        "#ISO": str(date_solicitate.get('Certificări ISO', 'N/A')),
        "#activitate_curenta": str(date_solicitate.get('Activitate curentă', 'N/A')),
        "#dotari_activitate_curenta": str(date_solicitate.get('Dotări pentru activitatea curentă', 'N/A')),
        "#info_ctr_implementare": str(date_solicitate.get('Informații despre contractul de implementare', 'N/A')),
        "#zonele_vizate_prioritar": str(date_solicitate.get('Zonele vizate prioritare', 'N/A')),
        "#utilaj_ghidare": str(date_solicitate.get('Utilaj de ghidare', 'N/A')),
        "#legaturi": str(date_solicitate.get('Legături', 'N/A')),
        "#rude": str(date_solicitate.get('Rude în cadrul firmei', 'N/A')),
        "#concluzie_CA": str(date_solicitate.get('Concluzie_CA', 'N/A')),
        "#caracteristici_tehnice": str(date_solicitate.get('Caracteristici tehnice relevante', 'N/A')),
        "#flux_tehnologic": str(date_solicitate.get('Flux tehnologic', 'N/A')),
        "#utilajeDNSH": str(date_solicitate.get('Utilaje DNSH', 'N/A')),
        "#descriere_utilaj_ghidare": str(date_solicitate.get('Descriere utilaj ghidare', 'N/A')),
        "#descriere_utilaj_reciclare": str(date_solicitate.get('Descriere utilaj reciclare', 'N/A')),
        "#contributia_proiectului_la_TJ": str(date_solicitate.get('Contribuția proiectului la tranziția justă', 'N/A')),
        "#strategii_materiale": str(date_solicitate.get('Strategii materiale', 'N/A')),
        "#strategii_reciclate": str(date_solicitate.get('Strategii materiale reciclate', 'N/A')),
        "#activitate": str(date_solicitate.get('Activitate specifică', 'N/A')),
        "#descriere_utilaj_reciclare": str(date_solicitate.get('Descriere utilaj de reciclare', 'N/A')),
        "#lucrari_inovatie": str(date_solicitate.get('Inovații în lucrări', 'N/A')),
        "#lucrari_caen": str(date_solicitate.get('Lucrări conform codurilor CAEN', 'N/A')),
        "#aDNSH": str(date_solicitate.get('Detalii DNSH - A', 'N/A')),
        "#cDNSH": str(date_solicitate.get('Detalii DNSH - C', 'N/A')),
        "#dDNSH": str(date_solicitate.get('Detalii DNSH - D', 'N/A')),
        "#materiale_locale": str(date_solicitate.get('Utilizarea materialelor locale', 'N/A')),
        "#PregatireaTeren": str(date_solicitate.get('Pregătirea terenului pentru lucrări', 'N/A')),
        "#ReciclareaMaterialelor": str(date_solicitate.get('Procesul de reciclare a materialelor', 'N/A')),
        "#clientiFirma": str(date_solicitate.get('Clienți principali ai firmei', 'N/A')),
        "#DacaTipInvest": str(date_solicitate.get('Tipul investiției planificate', 'N/A')),
        "#crearea": str(date_solicitate.get('Crearea de noi oportunități', 'N/A')),
        "#CompletDiversificare": str(date_solicitate.get('Diversificarea activităților firmei', 'N/A')),
        "#CompletExtinderea": str(date_solicitate.get('Extinderea capacității firmei', 'N/A')),
        "#CrestCreare": str(date_solicitate.get('Creșterea și crearea de noi activități', 'N/A')),
        "#CreareActivVizata": str(date_solicitate.get('Crearea de activități în domeniul vizat', 'N/A')),
        "#DezavantajeConcurentiale": str(date_solicitate.get('Identificarea dezavantajelor concurențiale', 'N/A')),
        "#30nrLocMunca": str(date_solicitate.get('Locuri Noi Create 30%', 'N/A')),
        "#20NrLocMunca": str(date_solicitate.get('Locuri Noi Create 20%', 'N/A')),
        "#zoneDN": str(date_solicitate.get('Zone vizate Prioritar', 'N/A')),
        "#Iso14001": str(date_solicitate.get('Certificat ISO 14001', 'N/A')),
        
        
        "#NAM20": str(situatie_angajati.get('Numar mediu angajati 2020', 'N/A')),
        "#NAM21": str(situatie_angajati.get('Numar mediu angajati 2021', 'N/A')),
        "#NAM22": str(situatie_angajati.get('Numar mediu angajati 2022', 'N/A')),   
        "#CPA20": str(capital_propriu.get('Capitalul propriu al actionarilor 2020', 'N/A')), 
        "#CPA21": str(capital_propriu.get('Capitalul propriu al actionarilor 2021', 'N/A')),
        "#CPA22": str(capital_propriu.get('Capitalul propriu al actionarilor 2022', 'N/A')),
        "#CA20": str(cifra_venit_rezultat.get('Cifra de afaceri 2020', 'N/A')),
        "#CA21": str(cifra_venit_rezultat.get('Cifra de afaceri 2021', 'N/A')),
        "#CA22": str(cifra_venit_rezultat.get('Cifra de afaceri 2022', 'N/A')),
        "#VT20": str(cifra_venit_rezultat.get('Venituri totale 2020', 'N/A')),
        "#VT21": str(cifra_venit_rezultat.get('Venituri totale 2021', 'N/A')),
        "#VT22": str(cifra_venit_rezultat.get('Venituri totale 2022', 'N/A')),     
        "#REX20": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2020', 'N/A')),
        "#REX21": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2021', 'N/A')), 
        "#REX22": str(cifra_venit_rezultat.get('Rezultat al exercitiului 2022', 'N/A')),
        "#MAXCA": str(cifra_venit_rezultat.get('Anul cu cea mai mare cifra de afaceri', 'N/A')),
        "#RSG20": str(rata_rent_grad.get('Rata solvabilitatii generale 2020', 'N/A')),
        "#RSG21": str(rata_rent_grad.get('Rata solvabilitatii generale 2021', 'N/A')), 
        "#RSG22": str(rata_rent_grad.get('Rata solvabilitatii generale 2022', 'N/A')),
        "#GITS20": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2020', 'N/A')),
        "#GITS21": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2021', 'N/A')),
        "#GITS22": str(rata_rent_grad.get('Gradul de indatorare pe termen scurt 2022', 'N/A')),
        "#ROA20": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2020', 'N/A')),
        "#ROA21": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2021', 'N/A')),
        "#ROA22": str(rata_rent_grad.get('Rentabilitatea activelor (ROA) 2022', 'N/A')),
        "#ROE20": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2020', 'N/A')),
        "#ROE21": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2021', 'N/A')),
        "#ROE22": str(rata_rent_grad.get('Rentabilitatea capitalului propriu (ROE) 2022', 'N/A')),
  
    }

    def inlocuieste_in_tabele(tabele, placeholders):
        for tabel in tabele:
            for row in tabel.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, value in placeholders.items():
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                    # Verifică dacă există tabele încastrate în celulă și aplică funcția recursiv
                    if cell.tables:
                        inlocuieste_in_tabele(cell.tables, placeholders)

    inlocuieste_in_tabele(template_doc.tables, placeholders)

    for paragraph in template_doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholders.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    modified_doc_path = "document_modificat.docx"
    template_doc.save(modified_doc_path)

    with open(modified_doc_path, "rb") as file:
        st.download_button(label="Descarcă Documentul Completat", data=file, file_name="document_modificat.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
