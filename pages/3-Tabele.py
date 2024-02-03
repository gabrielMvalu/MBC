import streamlit as st
import pandas as pd
import io
from docx import Document
from docx.shared import Pt 

st.header('Pregătirea datelor din P. FINANCIAR pentru completare tabel subcap 2.4')

uploaded_file = st.file_uploader("Încarcă documentul '*.xlsx' aici", type="xlsx", accept_multiple_files=False)
stop_text = "Total proiect"

def transforma_date(df):
    stop_index = df.index[df.iloc[:, 1].eq(stop_text)].tolist()
    df = df.iloc[3:stop_index[0]] if stop_index else df.iloc[3:]
    df = df[df.iloc[:, 1].notna() & (df.iloc[:, 1] != 0) & (df.iloc[:, 1] != '-')]

    nr_crt, um_list, cantitate_list, pret_unitar_list, valoare_totala_list, linie_bugetara_list, eligibil_neeligibil = [], [], [], [], [], [], []
    counter = 1

    for index, row in df.iterrows():
    item = row.iloc[1].strip().lower()
    if item in ["total active corporale", "total active necorporale"]:
        nr_crt.append(None)
        um_list.append(None)
        cantitate_list.append(None)
        pret_unitar_list.append(None)
        valoare_totala_list.append(None)
        linie_bugetara_list.append(None)
    else:
        nr_crt.append(counter)
        um_list.append("buc")

        cantitate = int(row.iloc[11]) if pd.notna(row.iloc[11]) else None
        cantitate_list.append(cantitate)

        pret_unitar_list.append(row.iloc[3])
        valoare_totala = row.iloc[3] * cantitate if cantitate is not None else None
        valoare_totala_list.append(valoare_totala)
        linie_bugetara_list.append(row.iloc[14])
        counter += 1

    for index, row in df.iterrows():
        val_6 = pd.to_numeric(row.iloc[6], errors='coerce')
        val_4 = pd.to_numeric(row.iloc[4], errors='coerce')
        if pd.isna(val_6) or pd.isna(val_4):
            eligibil_neeligibil.append("Data Missing")
        elif val_6 == 0 and val_4 != 0:
            eligibil_neeligibil.append(f"0 // {round(val_4, 2)}")
        elif val_6 == 0 and val_4 == 0:
            eligibil_neeligibil.append("0 // 0")
        elif val_6 < val_4:
            eligibil_neeligibil.append(f"{round(val_6, 2)} // {round(val_4 - val_6, 2)}")
        else:
            eligibil_neeligibil.append(f"{round(val_6, 2)} // {round(val_6 - val_4, 2)}")

    df_nou = pd.DataFrame({
        "Nr. crt.": nr_crt,
        "Denumirea lucrărilor / bunurilor/ serviciilor": df.iloc[:, 1],
        "UM": um_list,
        "Cantitate": cantitate_list,
        "Preţ unitar (fără TVA)": pret_unitar_list,
        "Valoare Totală (fără TVA)": valoare_totala_list,
        "Linie bugetară": linie_bugetara_list,
        "Eligibil/ neeligibil": eligibil_neeligibil,
        "Contribuie la criteriile de evaluare a,b,c,d": df.iloc[:, 15]
    })

    return df_nou

def df_to_word(document, df):
    table = document.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
    for j in range(df.shape[1]):
        cell = table.cell(0, j)
        cell.text = df.columns[j]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.cell(i + 1, j)
            val = df.iloc[i, j]
            cell.text = "" if pd.isna(val) else str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    return document

if st.button("Generează Tabel 1"):
    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file, sheet_name="P. FINANCIAR")
        tabel_1 = transforma_date(df)
        doc = Document()
        doc = df_to_word(doc, tabel_1)
        towrite = io.BytesIO()
        doc.save(towrite)
        towrite.seek(0)
        st.download_button(label="Descarcă Tabelul 1 ca Word",
                           data=towrite,
                           file_name="tabel_prelucrat.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.error("Te rog să încarci un fișier.")

