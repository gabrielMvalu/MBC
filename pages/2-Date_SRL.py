#pages/Date_srl.py
import streamlit as st
from docx import Document
import re

def extrage_informatii_firma(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    company_pattern1 = r"informațiile referitoare la\s*(.*?)\s*INFORMAȚII DE IDENTIFICARE"
    company_pattern2 = r"FURNIZARE INFORMAŢII\n\n(.*?)\n"
    firma_match = re.search(company_pattern1, full_text, re.IGNORECASE | re.DOTALL)
    if firma_match:
        firma = firma_match.group(1).strip()
    else:
        firma_match = re.search(company_pattern2, full_text, re.DOTALL)
        firma = firma_match.group(1).strip() if firma_match else "N/A"
  
       # Attempt to extract the registration number in its standard location or from the EUID line
    nr_ordine_patterns = [
        r"Număr de ordine în Registrul Comerțului:\s*([\w/]+)",
        r"EUID:\s*ROONRC\.([\w/]+)"
    ]
    
    # Check each pattern and break once a match is found
    nr_ordine = "N/A"
    for pattern in nr_ordine_patterns:
        nr_ordine_match = re.search(pattern, full_text)
        if nr_ordine_match:
            nr_ordine = nr_ordine_match.group(1)
            break

    
    cui_match = re.search(r"Cod unic de înregistrare: (\d+)", full_text)
    cui = cui_match.group(1) if cui_match else "N/A"
    data_infiintarii_match = re.search(r"atribuit în data de (\d+\.\d+\.\d+)", full_text)
    data_infiintarii = data_infiintarii_match.group(1) if data_infiintarii_match else "N/A"
    address_pattern = re.compile(r"Adresă sediu social: (.*?)(?=\n)")
    address_match = re.search(address_pattern, full_text)
    adresa = address_match.group(1) if address_match else "N/A"
    main_activity_pattern = r"Activitatea principală.*?Domeniul de activitate principal:.*?\n(.*?)(?:\n|;)"
    main_activity_match = re.search(main_activity_pattern, full_text, re.DOTALL)
    main_activity = main_activity_match.group(1).strip() if main_activity_match else "N/A"
    section_pattern = re.compile(r"SEDII SECUNDARE / PUNCTE DE LUCRU(.*?)SEDII SI/SAU ACTIVITATI AUTORIZATE", re.DOTALL)
    section_match = re.search(section_pattern, full_text)
    if section_match:
        section_text = section_match.group(1)
        secondary_address_pattern = re.compile(r"Adresă: (.*?)(?=\n)", re.DOTALL)
        adrese_secundare = re.findall(secondary_address_pattern, section_text)
    else:
        adrese_secundare = ["N/A"]
    data = {
        "Denumirea firmei": firma,
        "Numărul de ordine în Registrul Comerțului": nr_ordine,
        "Codul unic de înregistrare (CUI)": cui,
        "Data înființării": data_infiintarii,
        "Adresa sediului social": adresa,
        "Activitate principală": main_activity,
        "Adresa sediul secundar": adrese_secundare
    }
    return data

def extract_detailed_info_from_docx(doc):
    text = [p.text for p in doc.paragraphs]
    asociati = {}
    administratori = set()
    in_asociati_section = False
    in_persoane_imputernicite_section = False

    for i in range(len(text)):
        if "ASOCIAŢI PERSOANE FIZICE" in text[i]:
            in_asociati_section = True
            continue
        elif "REPREZENTANT acţionar/asociat/membru" in text[i]:
            in_asociati_section = False

        if in_asociati_section and "Calitate: " in text[i]:
            nume = text[i - 1].strip()
            j = i + 1
            while "Cota de participare la beneficii şi pierderi: " not in text[j] and j < len(text) - 1:
                j += 1
            if j < len(text):
                cota = text[j].split(":")[1].strip()
                asociati[nume] = cota

        if "Persoane împuternicite (PERSOANE FIZICE)" in text[i]:
            in_persoane_imputernicite_section = True
            continue
        elif "Persoane împuternicite (PERSOANE JURIDICE)" in text[i]:
            in_persoane_imputernicite_section = False

        if in_persoane_imputernicite_section and "Calitate: " in text[i]:
            nume_admin = text[i - 1].strip()
            administratori.add(nume_admin)

    output_asociati = []
    for nume, cota in asociati.items():
        info = f"{nume} – asociat cu cota de participare la beneficii și pierderi {cota}"
        if nume in administratori:
            info += " și administrator"
        output_asociati.append(info)

    nume_administrator = ', '.join(administratori)  # Gestionarea cazului cu mai mulți administratori

    return output_asociati, nume_administrator

def extract_situatie_angajati(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    angajati_pattern_2020 = r"SITUAŢIA FINANCIARĂ PE ANUL 2020.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2020 = re.search(angajati_pattern_2020, full_text, re.DOTALL)
    nrang20 = angajati_match_2020.group(1) if angajati_match_2020 else "N/A"
    angajati_pattern_2021 = r"SITUAŢIA FINANCIARĂ PE ANUL 2021.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2021 = re.search(angajati_pattern_2021, full_text, re.DOTALL)
    nrang21 = angajati_match_2021.group(1) if angajati_match_2021 else "N/A"
    angajati_pattern_2022 = r"SITUAŢIA FINANCIARĂ PE ANUL 2022.*?(?:Numar|Număr) mediu de salari(?:aţi|ati): (\d+)"
    angajati_match_2022 = re.search(angajati_pattern_2022, full_text, re.DOTALL)
    nrang22 = angajati_match_2022.group(1) if angajati_match_2022 else "N/A"
    data_angajati = {
        "Numar mediu angajati 2020": nrang20,
        "Numar mediu angajati 2021": nrang21,
        "Numar mediu angajati 2022": nrang22,
    }
    return data_angajati


#def extract_caen_codes(full_text):
#    start_marker = "SEDII SI/SAU ACTIVITATI AUTORIZATE"
#    end_marker = "CONCORDAT PREVENTIV"
#    caen_section_pattern = re.compile(rf"{start_marker}(.*?){end_marker}", re.DOTALL)
#    caen_section_match = re.search(caen_section_pattern, full_text)
    
#    if caen_section_match:
#        caen_section_text = caen_section_match.group(1)
#        caen_code_pattern = re.compile(r"(\d{4}) - (.*?)\n")
#        caen_codes = re.findall(caen_code_pattern, caen_section_text)
#        return caen_codes
#    else:
#        return []

#Modicicat in urmatoarea functie conform noilor cerinte de afisare!!! 12 feb 2023




def extrage_coduri_caen(doc):
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    start_marker = "SEDII SI/SAU ACTIVITATI AUTORIZATE"
    end_marker = "Denumire: Punct de lucru"
    pattern = fr"(?s){start_marker}(.*?){end_marker}"

    results = []
    matches = re.findall(pattern, full_text)
    for match in matches:
        # Verificăm dacă secțiunea conține textul specificat pentru activitățile neautorizate
        if "Nu se desfăşoară activităţile prevăzute în actul constitutiv sau modificator" not in match:
            # Extragem informațiile despre tipul de activitate autorizată și codurile CAEN
            tip_activitate_pattern = r"Tip activitate autorizată: terţi\n(?:Conform declaraţiei.*?\n)?(?:Activităţi desfăşurate în afara sediului social şi a sediilor secundare \(CAEN REV\. 2\):\s*)?((?:\d{4} - .+?(?:\n|$))+)"
            tip_activitate_match = re.search(tip_activitate_pattern, match, re.DOTALL)
            if tip_activitate_match:
                tip_activitate_info = tip_activitate_match.group(1).strip()
                # Eliminăm tot ce urmează după ultimul cod CAEN, inclusiv "Data certificatului constatator"
                tip_activitate_info = re.sub(r"\nData certificatului.*$", "", tip_activitate_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre tipul de activitate autorizată cu codurile CAEN
                combined_info = f"Tip activitate autorizată: terţi\nActivităţi desfăşurate în afara sediului social şi a sediilor secundare:\n{tip_activitate_info}"
                results.append(combined_info)

            # Extragem întreaga adresă a sediului și activitățile la sediu
            sediu_info_match = re.search(r"(Sediul (social|secundar|terţ) din:.+?)(?=Tip sediu:)", match, re.DOTALL)
            sediu_info = sediu_info_match.group(1).strip() if sediu_info_match else ""

            activitati_pattern = r"Activităţi la sediu:\s*((?:\d{4} - .+?(?:\n|$))+)"
            activitati_match = re.search(activitati_pattern, match, re.DOTALL)
            if activitati_match:
                activitati_info = activitati_match.group(1).strip()
                activitati_info = re.sub(r"\nData certificatului.*$", "", activitati_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre sediu cu activitățile la sediu
                combined_info = f"{sediu_info}\nActivităţi la sediu:{activitati_info}"
                results.append(combined_info)
        else:
            tip_activitate_pattern = r"Tip activitate autorizată: terţi\n(?:Conform declaraţiei.*?\n)?(?:Activităţi desfăşurate în afara sediului social şi a sediilor secundare \(CAEN REV\. 2\):\s*)?((?:\d{4} - .+?(?:\n|$))+)"
            tip_activitate_match = re.search(tip_activitate_pattern, match, re.DOTALL)
            if tip_activitate_match:
                tip_activitate_info = tip_activitate_match.group(1).strip()
                # Eliminăm tot ce urmează după ultimul cod CAEN, inclusiv "Data certificatului constatator"
                tip_activitate_info = re.sub(r"\nData certificatului.*$", "", tip_activitate_info, flags=re.MULTILINE).strip()
                # Combinăm informațiile despre tipul de activitate autorizată cu codurile CAEN
                combined_info = f"Tip activitate autorizată: terţi\nActivităţi desfăşurate în afara sediului social şi a sediilor secundare:\n{tip_activitate_info}"
                results.append(combined_info)


    return results




# Încărcarea și procesarea documentului în Streamlit
st.header(':blue[Încărcare Document Registrul Comerțului]',divider='rainbow')

uploaded_file = st.file_uploader("Trageți fișierul aici sau faceți clic pentru a încărca un document", type=["docx"])

if uploaded_file is not None:
    doc = Document(uploaded_file)
    general_data = extrage_informatii_firma(doc)
    detailed_info, admins = extract_detailed_info_from_docx(doc)
    angajati_data = extract_situatie_angajati(doc)
    sedii_si_activitati = extrage_coduri_caen(doc)
    # anulat in urma cerintei din 12 feb 2023  caen_codes = extract_caen_codes("\n".join([p.text for p in doc.paragraphs]))



    # Afișarea datelor în format JSON
    st.json({
        "Date Generale": general_data,
        "Informații Detaliate": {"Asociați": detailed_info, "Administratori": admins},
        "Situație Angajati": angajati_data,
        "Sedii si activititati plus CAEN": sedii_si_activitati
    })

    # Afișarea datelor în DataFrame-uri pentru vizualizare
 #   st.write("Date Generale:")
 #   st.dataframe(pd.DataFrame([general_data]))

    # Convertirea listei de administratori într-o listă de dicționare
  #  admins_dict = [{"Administratori": admin} for admin in admins]
  #  # Crearea DataFrame-ului din lista de dicționare
  #  admins_df = pd.DataFrame(admins_dict)
    # Afișarea DataFrame-ului în Streamlit
  #  st.write("Administratori:")
  #  st.dataframe(admins_df)


  #  st.write("Situație Financiară:")
  #  st.dataframe(pd.DataFrame([angajati_data]))

  #  st.write("Sedii si activititati plus CAEN:")
  #  st.dataframe(pd.DataFrame(sedii_si_activitati)

    # După extragere, salvează datele în session_state
  #  st.session_state['date_generale'] = general_data
  #  st.session_state['date_detaliat'] = {"Asociați": detailed_info, "Administratori": admins}
  #  st.session_state['situatie_angajati'] = angajati_data
  #  st.session_state['coduri_caen'] = sedii_si_activitati
