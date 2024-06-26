# pages/Bilant & Analiza.py
# pages/Bilant & Analiza.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import altair as alt
from altair import Scale, Color
import io
from docx import Document

st.set_page_config(layout="wide")


def extrage_date_bilant(df):
    cpa20 = f"{df.iloc[99, 1]:.2f}"
    cpa21 = f"{df.iloc[99, 2]:.2f}"
    cpa22 = f"{df.iloc[99, 3]:.2f}"
    
    data = {
        "Capitalul propriu al actionarilor 2020": cpa20, 
        "Capitalul propriu al actionarilor 2021": cpa21,
        "Capitalul propriu al actionarilor 2022": cpa22
    }
    return data

def extrage_date_contpp(df1):
    ca20 = f"{df1.iloc[4, 1]:.2f}"
    ca21 = f"{df1.iloc[4, 2]:.2f}"
    ca22 = f"{df1.iloc[4, 3]:.2f}"
    vt20 = f"{df1.iloc[55, 1]:.2f}"
    vt21 = f"{df1.iloc[55, 2]:.2f}"
    vt22 = f"{df1.iloc[55, 3]:.2f}"

    if df1.iloc[4, 6] != 0:
        pc = f"{(df1.iloc[4, 6] / df1.iloc[4, 3] * 100) - 100:.2f}%"
    else:
        pc = "N/A" 
    
    if df1.iloc[4, 1] > df1.iloc[4, 2] and df1.iloc[4, 1] > df1.iloc[4, 3]:
        camax = 2020
    elif df1.iloc[4, 2] > df1.iloc[4, 1] and df1.iloc[4, 2] > df1.iloc[4, 3]:
        camax = 2021
    else:
        camax = 2022
    
    re20 = f"{df1.iloc[66, 1]:.2f}" if df1.iloc[66, 1] > 0 else f"{df1.iloc[67, 1]:.2f}"
    re21 = f"{df1.iloc[66, 2]:.2f}" if df1.iloc[66, 2] > 0 else f"{df1.iloc[67, 2]:.2f}"
    re22 = f"{df1.iloc[66, 3]:.2f}" if df1.iloc[66, 3] > 0 else f"{df1.iloc[67, 3]:.2f}"
    
    data = {
        "Cifra de afaceri 2020": ca20, 
        "Cifra de afaceri 2021": ca21,
        "Cifra de afaceri 2022": ca22,
        "Venituri totale 2020": vt20, 
        "Venituri totale 2021": vt21,
        "Venituri totale 2022": vt22,
        "Rezultat al exercitiului 2020": re20, 
        "Rezultat al exercitiului 2021": re21,
        "Rezultat al exercitiului 2022": re22,  
        "Anul cu cea mai mare cifra de afaceri": camax, 
        "Procent crestere": pc,
    }
    return data

def extrage_indicatori_financiari(df2):
    rs20 = f"{df2.iloc[89, 1]:.2f}"  
    rs21 = f"{df2.iloc[89, 2]:.2f}"
    rs22 = f"{df2.iloc[89, 3]:.2f}"
    gdi20 = f"{df2.iloc[95, 1]:.0%}"
    gdi21 = f"{df2.iloc[95, 2]:.0%}"
    gdi22 = f"{df2.iloc[95, 3]:.0%}"
    roa20 = f"{df2.iloc[43, 1]:.0%}"
    roa21 = f"{df2.iloc[43, 2]:.0%}"
    roa22 = f"{df2.iloc[43, 3]:.0%}"
    roe20 = f"{df2.iloc[47, 1]:.0%}"
    roe21 = f"{df2.iloc[47, 2]:.0%}"
    roe22 = f"{df2.iloc[47, 3]:.0%}"

    data = {
        "Rata solvabilitatii generale 2020": rs20, 
        "Rata solvabilitatii generale 2021": rs21,
        "Rata solvabilitatii generale 2022": rs22,
        "Gradul de indatorare pe termen scurt 2020": gdi20, 
        "Gradul de indatorare pe termen scurt 2021": gdi21,
        "Gradul de indatorare pe termen scurt 2022": gdi22,
        "Rentabilitatea activelor (ROA) 2020": roa20, 
        "Rentabilitatea activelor (ROA) 2021": roa21,
        "Rentabilitatea activelor (ROA) 2022": roa22,
        "Rentabilitatea capitalului propriu (ROE) 2020": roe20, 
        "Rentabilitatea capitalului propriu (ROE) 2021": roe21,
        "Rentabilitatea capitalului propriu (ROE) 2022": roe22,
    }

    return data
st.header(':blue[Adaugati Analiză, Bilanț, Cont de Profit și Pierdere]', divider='rainbow')

uploaded_file = st.file_uploader("Adăugați fișierul aici sau faceți click pentru a încărca", type=["xlsx"])

#on = st.toggle('Activate feature')




def create_word_document(data_bilant, data_contpp, data_analiza):
    doc = Document()
    doc.add_heading('Raport Analiza', 0)

    for data, title in zip([data_bilant, data_contpp, data_analiza], ['Bilanț', 'Cont de Profit și Pierdere', 'Analiză']):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=len(data) + 1, cols=2)  # 2 coloane: Numele indicatorului și Valoarea
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Indicator'
        hdr_cells[1].text = 'Valoare'
        
        for i, (key, value) in enumerate(data.items(), start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        doc.add_paragraph()

    return doc




if uploaded_file is not None:
    df = pd.read_excel(uploaded_file, sheet_name='1-Bilant')
    df1 = pd.read_excel(uploaded_file, sheet_name='2-ContPP')
    df2 = pd.read_excel(uploaded_file, sheet_name='1D-Analiza_fin_indicatori')    
    data_bilant = extrage_date_bilant(df)
    data_contpp = extrage_date_contpp(df1)
    data_analiza = extrage_indicatori_financiari(df2) 

    # Accesează valorile din dicționarul 'data_contpp'
   # ca22 = data_contpp["Cifra de afaceri 2022"]
   # pc = data_contpp["Procent crestere"]

    # if on:
       # Utilizează valorile în 'st.metric'
    #   st.metric(label="Procentul de creștere CA", value=pc, delta=ca22, 
    #   delta_color="off", help='Procentul trebuie sa fie cat mai mic, deoarece este indicator la proiect si trebuie sa il respecte')


    if st.button('Descărcați Raportul'):
        doc = create_word_document(data_bilant, data_contpp, data_analiza)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(label="Descărcați Raportul Word",
                           data=buffer,
                           file_name="Raport_Financiar.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    
   

    # Pregătirea datelor pentru grafice
    df_cifra_afaceri = pd.DataFrame({
        'An': ['2020', '2021', '2022'],
        'Cifra de afaceri': [float(data_contpp["Cifra de afaceri 2020"]), float(data_contpp["Cifra de afaceri 2021"]), float(data_contpp["Cifra de afaceri 2022"])],
        'Venituri totale': [float(data_contpp["Venituri totale 2020"]), float(data_contpp["Venituri totale 2021"]), float(data_contpp["Venituri totale 2022"])]
    })
    
    df_solvabilitate = pd.DataFrame({
        'An': ['2020', '2021', '2022'],
        'Rata solvabilității generale': [float(data_analiza["Rata solvabilitatii generale 2020"]), float(data_analiza["Rata solvabilitatii generale 2021"]), float(data_analiza["Rata solvabilitatii generale 2022"])]
    })
    
    culori_customizate = Scale(domain=['2020', '2021', '2022'], 
                               range=['#f7e928', '#28f7e9', '#fa1b1b'])
    
    col1, col2 = st.columns(2)
    
    with col2:
        # Crearea și afișarea primului grafic în col1
        st.write('## Evoluția Cifrei de Afaceri și a Veniturilor Totale')
        grafic_cifra_afaceri = alt.Chart(df_cifra_afaceri.melt('An', var_name='Categorie', value_name='Valoare')).mark_line(point=True).encode(
            x='An:N',
            y='Valoare:Q',
            color='Categorie:N',
            tooltip=['An', 'Categorie', 'Valoare']
        ).interactive()
        st.altair_chart(grafic_cifra_afaceri, use_container_width=True)
    
    
    with col1:
        # Crearea și afișarea celui de-al doilea grafic în col2
        st.write("## Rata Solvabilității Generale pe Ani")
        grafic_solvabilitate = alt.Chart(df_solvabilitate).mark_bar().encode(
            x='An:N',
            y='Rata solvabilității generale:Q',
            color=Color('An:N', scale=culori_customizate),
            tooltip=['An', 'Rata solvabilității generale']
        ).interactive()
        st.altair_chart(grafic_solvabilitate, use_container_width=True)
    
    
    col3, col4 = st.columns(2)
    with col3:
         st.json({"Datele din bilant sunt:": data_bilant})
         st.json({"Datele din contPP sunt:": data_contpp})    
         st.json({"Datele din analiza sunt:": data_analiza})
          
    with col4:
        st.write("Vizualizare Bilant:")
        st.dataframe(pd.DataFrame([data_bilant]))
        st.dataframe(pd.DataFrame([data_contpp])) 
        st.dataframe(pd.DataFrame([data_analiza])) 



